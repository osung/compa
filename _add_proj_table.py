# -*- coding: utf-8 -*-
"""각 TopN 과제 상세 블록에 7개 필드 세로형 정보표 삽입(제목 아래). 기존 '수행기관:' 줄 제거."""
import re, sys, json
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

SRC = sys.argv[1] if len(sys.argv) > 1 else "COMPA_최종Top5_보고서.docx"
DST = sys.argv[2] if len(sys.argv) > 2 else SRC
SCRATCH = "/private/tmp/claude-501/-Users-osung-work-compa/d6ed121c-12e4-45b4-b2fb-535b7554627c/scratchpad"

def norm(s): return re.sub(r"\s+", "", str(s)).strip()

data = json.load(open(f"{SCRATCH}/pid_fields.json"))
name2pid = {tuple(k.split("||")): v for k, v in data["name2pid"].items()}  # (기업,수요,과제명)->pid
FIELDS = data["fields"]

def fmt_date(s):
    s = re.sub(r"\.0$", "", str(s)).strip()
    return f"{s[:4]}.{s[4:6]}.{s[6:8]}" if re.fullmatch(r"\d{8}", s) else s

def fmt_cost(s):
    s = re.sub(r"\.0$", "", str(s)).strip()
    try:
        w = float(s)
    except ValueError:
        return s or "-"
    if w >= 1e8:
        return f"{w/1e8:.1f}억원"
    if w >= 1e4:
        return f"{w/1e4:,.0f}만원"
    return f"{w:,.0f}원"

def rows_for(pid, org_fallback):
    f = FIELDS.get(pid, {})
    s, e = fmt_date(f.get("연구시작일", "")), fmt_date(f.get("연구종료일", ""))
    period = f"{s} ~ {e}" if s and e else (s or e or "-")
    dae, jung = f.get("표준분류대", ""), f.get("표준분류중", "")
    cls = f"{dae} > {jung}" if dae and jung else (jung or dae or "-")
    return [
        ("과제고유번호", f.get("과제고유번호", pid) or "-"),
        ("과제수행기간", period),
        ("과학기술표준분류", cls),
        ("연구개발단계", f.get("연구개발단계", "") or "-"),
        ("총연구비", fmt_cost(f.get("총연구비", ""))),
        ("과제수행기관", f.get("과제수행기관명", "") or org_fallback or "-"),
        ("연구수행주체", f.get("연구수행주체", "") or "-"),
    ]

W_LABEL, W_VALUE = 1900, 4300

def style_table(tbl):
    tblPr = tbl._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        x = OxmlElement(f"w:{e}")
        x.set(qn("w:val"), "single"); x.set(qn("w:sz"), "4")
        x.set(qn("w:space"), "0"); x.set(qn("w:color"), "808080")
        b.append(x)
    tblPr.append(b)
    lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"), "fixed"); tblPr.append(lay)
    w = OxmlElement("w:tblW"); w.set(qn("w:w"), str(W_LABEL + W_VALUE)); w.set(qn("w:type"), "dxa"); tblPr.append(w)
    ind = OxmlElement("w:tblInd"); ind.set(qn("w:w"), "240"); ind.set(qn("w:type"), "dxa"); tblPr.append(ind)

def set_w(cell, w):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
    tcW.set(qn("w:w"), str(w)); tcW.set(qn("w:type"), "dxa")

def fill(cell, text, label=False):
    p = cell.paragraphs[0]; p.text = ""
    r = p.add_run(text); r.font.size = Pt(9); r.bold = label
    if label:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F2F2F2")
        tcPr.append(shd)

d = Document(SRC)
# 본문 순회로 각 TopN 제목에 (기업명, 수요기술명) 맥락 부여 — 동일 과제명 충돌 방지
from docx.text.paragraph import Paragraph
from docx.table import Table
heads = []; cur_comp = cur_dem = None
for ch in list(d.element.body):
    if ch.tag == qn("w:p"):
        pp = Paragraph(ch, d); tt = pp.text.strip(); ss = pp.style.name if pp.style else ""
        if ss == "Heading 2" and tt.startswith("[수요"):
            cur_dem = re.sub(r"^\[수요\s*\d+\]\s*", "", tt).strip(); cur_comp = None
        elif ss == "Heading 3" and re.match(r"Top\d", tt):
            heads.append((pp, cur_comp, cur_dem))
    elif ch.tag == qn("w:tbl"):
        tbx = Table(ch, d)
        if tbx.rows[0].cells[0].text.strip() == "기업명" and cur_comp is None:
            cur_comp = tbx.rows[0].cells[1].text.strip()
n = miss = 0
for h, comp, dem in heads:
    m = re.match(r"Top\d+\.\s*(.*)", h.text.strip())
    title = re.sub(r"\s*\(출처[:：].*?\)\s*$", "", m.group(1)).strip()
    # 다음 형제에서 수행기관 줄 찾기
    org = ""; org_el = None
    el = h._p.getnext()
    steps = 0
    while el is not None and steps < 4:
        if el.tag == qn("w:p"):
            txt = "".join(t.text or "" for t in el.findall(".//" + qn("w:t")))
            if txt.strip().startswith("수행기관:"):
                org = txt.split("수행기관:", 1)[1].strip(); org_el = el; break
        el = el.getnext(); steps += 1
    pid = name2pid.get((norm(comp), norm(dem), norm(title)))
    if not pid:
        miss += 1; continue
    tbl = d.add_table(rows=7, cols=2)
    style_table(tbl)
    for ri, (lab, val) in enumerate(rows_for(pid, org)):
        lc, vc = tbl.rows[ri].cells
        set_w(lc, W_LABEL); set_w(vc, W_VALUE)
        fill(lc, lab, label=True); fill(vc, val)
    h._p.addnext(tbl._tbl)               # 제목 바로 뒤에 표
    if org_el is not None:               # 중복 '수행기관:' 줄 제거
        org_el.getparent().remove(org_el)
    n += 1

d.save(DST)
print(f"정보표 삽입: {n}개 | 매핑 실패: {miss}")
print("saved:", DST)
