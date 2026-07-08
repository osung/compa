# -*- coding: utf-8 -*-
"""과제 정보표 pid 충돌 오류(5행) 정정: 올바른 pid의 7필드 값으로 셀 재작성."""
import json, re, sys
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

def norm(s): return re.sub(r"\s+", "", str(s)).strip()
SCRATCH = "/private/tmp/claude-501/-Users-osung-work-compa/d6ed121c-12e4-45b4-b2fb-535b7554627c/scratchpad"
SRC = sys.argv[1] if len(sys.argv) > 1 else "COMPA_최종Top5_보고서.docx"
DST = sys.argv[2] if len(sys.argv) > 2 else SRC

FIELDS = json.load(open(f"{SCRATCH}/pid_fields.json"))["fields"]
targets = {tuple(k.split("||")): p for k, p in json.load(open(f"{SCRATCH}/fix_pids.json")).items()}

def fmt_date(s):
    s = re.sub(r"\.0$", "", str(s)).strip()
    return f"{s[:4]}.{s[4:6]}.{s[6:8]}" if re.fullmatch(r"\d{8}", s) else s
def fmt_cost(s):
    s = re.sub(r"\.0$", "", str(s)).strip()
    try: w = float(s)
    except ValueError: return s or "-"
    if w >= 1e8: return f"{w/1e8:.1f}억원"
    if w >= 1e4: return f"{w/1e4:,.0f}만원"
    return f"{w:,.0f}원"

def values_for(pid):
    f = FIELDS[pid]
    s, e = fmt_date(f["연구시작일"]), fmt_date(f["연구종료일"])
    period = f"{s} ~ {e}" if s and e else (s or e or "-")
    dae, jung = f.get("표준분류대", ""), f.get("표준분류중", "")
    cls = f"{dae} > {jung}" if dae and jung else (jung or dae or "-")
    return {"과제고유번호": pid, "과제수행기간": period, "과학기술표준분류": cls,
            "연구개발단계": f.get("연구개발단계", "") or "-", "총연구비": fmt_cost(f["총연구비"]),
            "과제수행기관": f.get("과제수행기관명", "") or "-", "연구수행주체": f.get("연구수행주체", "") or "-"}

def set_val(cell, text):
    p = cell.paragraphs[0]; runs = p.runs
    if not runs: p.add_run(text); return
    tgt = max(runs, key=lambda r: len(r.text)) if any(r.text for r in runs) else runs[-1]
    for r in runs: r.text = text if r is tgt else ""

d = Document(SRC)
cur = comp = lastname = None; fixed = 0
for ch in list(d.element.body):
    if ch.tag == qn("w:p"):
        p = Paragraph(ch, d); t = p.text.strip(); st = p.style.name if p.style else ""
        if st == "Heading 2" and t.startswith("[수요"): cur = re.sub(r"^\[수요\s*\d+\]\s*", "", t).strip(); comp = None
        elif st == "Heading 3" and re.match(r"Top\d", t): lastname = re.sub(r"\s*\(출처.*?\)\s*$", "", re.match(r"Top\d+\.\s*(.*)", t).group(1)).strip()
    elif ch.tag == qn("w:tbl"):
        tb = Table(ch, d); h0 = tb.rows[0].cells[0].text.strip()
        if h0 == "기업명" and comp is None: comp = tb.rows[0].cells[1].text.strip()
        elif h0 == "과제고유번호":
            k = (norm(comp), norm(cur), norm(lastname))
            if k in targets:
                vals = values_for(targets[k])
                for row in tb.rows:
                    lab = row.cells[0].text.strip()
                    if lab in vals: set_val(row.cells[1], vals[lab])
                fixed += 1
d.save(DST)
print(f"정보표 pid 정정: {fixed}건")
print("saved:", DST)
