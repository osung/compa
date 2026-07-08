# -*- coding: utf-8 -*-
"""양쪽 맞춤(justify) 적용:
 1) 본문 상세 매칭 근거 4섹션 문단([연관성]/[수요기술 사양 적합성]/[추천 과제의 우수성]/[유사 사례 및 실적])
 2) 수요 정보 표의 '수요기술 내용'·'수요기술 사양' 값 셀
 3) 최종 추천 Top5 표의 '매칭 근거' 열 데이터 셀."""
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH as A

SRC = sys.argv[1] if len(sys.argv) > 1 else "COMPA_최종Top5_보고서.docx"
DST = sys.argv[2] if len(sys.argv) > 2 else SRC

SECS = ("[연관성]", "[수요기술 사양 적합성]", "[추천 과제의 우수성]", "[유사 사례 및 실적]")
VALUE_ROWS = {"수요기술 내용", "수요기술 사양"}

d = Document(SRC)

# 1) 본문 4섹션 문단
n_sec = 0
for p in d.paragraphs:
    t = p.text.strip()
    if any(t.startswith(s) for s in SECS):
        p.alignment = A.JUSTIFY
        n_sec += 1

# 2)/3) 표
n_demand = n_reason = 0
for tb in d.tables:
    hdr = tuple(c.text.strip() for c in tb.rows[0].cells)
    if hdr[:1] == ("순위",) and "매칭 근거" in hdr:      # Top5 표
        mi = hdr.index("매칭 근거")
        for row in tb.rows[1:]:
            for para in row.cells[mi].paragraphs:
                para.alignment = A.JUSTIFY
            n_reason += 1
        continue
    # 수요 정보 표: 라벨(cell0)이 수요기술 내용/사양인 행의 값(cell1)
    for row in tb.rows:
        cells = row.cells
        if len(cells) >= 2 and cells[0].text.strip() in VALUE_ROWS:
            for para in cells[1].paragraphs:
                para.alignment = A.JUSTIFY
            n_demand += 1

d.save(DST)
print(f"상세근거 섹션 문단:{n_sec}  수요내용/사양 값셀:{n_demand}  매칭근거 셀:{n_reason}")
print("saved:", DST)
