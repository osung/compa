# -*- coding: utf-8 -*-
"""45건 재생성 입력 조립 + 27건 기존근거 복구. 모델 없이 데이터만 검증."""
import pickle, glob, json, re
import pandas as pd
from docx import Document

def norm(s): return re.sub(r"\s+", "", str(s)).strip()

SCRATCH = "/private/tmp/claude-501/-Users-osung-work-compa/d6ed121c-12e4-45b4-b2fb-535b7554627c/scratchpad"

# ---------- 1) 기존근거 조회표 (과제명 -> (기업명, 짧은근거, 상세)) ----------
by_name = {}   # norm(과제명) -> (기업명, short, detail)  (첫 등장 우선)
def add(name, comp, short, detail):
    k = norm(name)
    if k and k not in by_name:
        by_name[k] = (str(comp or ""), str(short or ""), str(detail or ""))

for f in glob.glob("COMPA_*_최종추천.pkl"):
    if "전체" in f: continue
    df = pickle.load(open(f, "rb"))
    for r in df.to_dict("records"):
        add(r.get("과제명", ""), r.get("기업명", ""), r.get("LLM판단근거", ""), r.get("추천근거_상세", ""))
j = json.load(open("COMPA_통합best.json"))
for k, e in j.items():
    for t in e.get("top5", []):
        add(t.get("과제명", ""), e.get("기업명", ""), t.get("판단근거", ""), t.get("추천근거_상세", ""))

# ---------- 2) docx 72건(상세근거 없는 Top) 추출 ----------
d = Document("COMPA_최종Top5_보고서.docx")
paras = [p.text.strip() for p in d.paragraphs]
styles = [p.style.name for p in d.paragraphs]
n = len(paras)
cur_comp = cur_dem = None
missing = []  # dict per Top without 상세 매칭 근거
for i, (t, s) in enumerate(zip(paras, styles)):
    if s == "Heading 2" and t.startswith("[수요"):
        cur_dem = re.sub(r"^\[수요\s*\d+\]\s*", "", t).strip()
    if s == "Normal" and t.startswith("기업명:"):
        cur_comp = t.split("기업명:", 1)[1].strip()
    if s == "Heading 3" and re.match(r"Top\d", t):
        j2 = i + 1; block = []
        while j2 < n and styles[j2] not in ("Heading 1", "Heading 2", "Heading 3"):
            block.append(paras[j2]); j2 += 1
        org = next((b.split("수행기관:", 1)[1].strip() for b in block if b.startswith("수행기관:")), "")
        if not any(b.startswith("상세 매칭 근거") for b in block):
            m = re.match(r"Top\d+\.\s*(.*)", t)
            name = re.sub(r"\s*\(출처[:：].*?\)\s*$", "", m.group(1)).strip()
            missing.append({"기업명": cur_comp, "수요기술명": cur_dem, "과제명": name, "수행기관": org})

print("docx 상세근거 없는 Top:", len(missing))

# ---------- 3) 27 복구 vs 45 생성 분리 ----------
recover, regen = [], []
for m in missing:
    v = by_name.get(norm(m["과제명"]))
    if v and v[2].strip():
        recover.append({**m, "판단근거": v[1], "추천근거_상세": v[2]})
    else:
        regen.append(m)
print("복구(27):", len(recover), " / 생성(45):", len(regen))

# ---------- 4) 후보풀에서 pid·수요필드·과제설명문 조인 (기업명,수요기술명,과제명) ----------
hp = pickle.load(open("COMPA_후보풀.pkl", "rb"))
hp_idx = {}
for r in hp.to_dict("records"):
    key = (norm(r["기업명"]), norm(r["수요기술명"]), norm(r["과제명"]))
    hp_idx.setdefault(key, r)  # first

# 키워드: 번호 -> keywords list  (담당자 pkl 의 키워드 컬럼)
kw_by_num = {}
dem_by_num = {}  # 번호 -> (기업명, 수요기술명) for sanity
for f in glob.glob("COMPA_*_최종추천.pkl"):
    if "전체" in f: continue
    df = pickle.load(open(f, "rb"))
    for r in df.to_dict("records"):
        num = str(r.get("번호"))
        kws = str(r.get("키워드") or "")
        kw_by_num[num] = [x for x in kws.split(";") if x.strip()]

regen_full, issues = [], []
for m in regen:
    key = (norm(m["기업명"]), norm(m["수요기술명"]), norm(m["과제명"]))
    r = hp_idx.get(key)
    if r is None:
        issues.append(("no후보풀", m["과제명"][:30])); continue
    num = str(r["번호"])
    kws = kw_by_num.get(num, [])
    m2 = dict(m)
    m2["번호"] = num
    m2["pid"] = str(r["과제고유번호"])
    m2["수요기술 내용"] = str(r.get("수요기술_내용") or "")
    m2["수요기술 사양"] = str(r.get("수요기술_사양") or "")
    m2["과제설명문_fallback"] = str(r.get("과제설명문") or "")
    m2["출처"] = str(r.get("출처") or "")
    m2["keywords"] = kws
    if not kws:
        issues.append(("no키워드", m["과제명"][:30]))
    regen_full.append(m2)
print("regen joined:", len(regen_full), "/", len(regen))
print("join issues:", issues[:20])

json.dump({"recover": recover, "regen": regen_full},
          open(f"{SCRATCH}/regen_targets.json", "w"), ensure_ascii=False, indent=1)
print("saved regen_targets.json ; recover=%d regen=%d" % (len(recover), len(regen_full)))
