# -*- coding: utf-8 -*-
"""Top5 표 서식: 순위·수행기관·과제수행년도 가운데 정렬, 제목행 전체 가운데 정렬,
과제수행년도는 '시작년도~' / '종료년도' 두 줄로 표시."""
import copy, sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = sys.argv[1] if len(sys.argv) > 1 else "COMPA_최종Top5_보고서.docx"
DST = sys.argv[2] if len(sys.argv) > 2 else SRC

CENTER_COLS = {0, 2, 3}   # 순위, 수행기관, 과제수행년도

def center(cell):
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def year_two_lines(cell):
    p = cell.paragraphs[0]
    runs = p.runs
    if not runs:
        return
    txt = "".join(r.text for r in runs).strip()
    if "~" not in txt:
        return  # 단일 연도는 그대로(한 줄)
    a, b = txt.split("~", 1)
    content = max(runs, key=lambda r: len(r.text)) if any(r.text for r in runs) else runs[-1]
    for r in runs:
        if r is not content:
            r.text = ""
    content.text = a + "~"
    new_r = OxmlElement("w:r")
    rPr = content._r.find(qn("w:rPr"))
    if rPr is not None:
        new_r.append(copy.deepcopy(rPr))
    new_r.append(OxmlElement("w:br"))
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = b
    new_r.append(t)
    content._r.addnext(new_r)

d = Document(SRC)
n = 0
for tb in d.tables:
    hdr = tuple(c.text.strip() for c in tb.rows[0].cells)
    if hdr[:1] != ("순위",) or "과제수행년도" not in hdr:
        continue
    yi = hdr.index("과제수행년도")
    # 제목행 전체 가운데 정렬
    for c in tb.rows[0].cells:
        center(c)
    # 데이터행
    for row in tb.rows[1:]:
        cells = row.cells
        for ci in CENTER_COLS:
            center(cells[ci])
        year_two_lines(cells[yi])
    n += 1

d.save(DST)
print(f"서식 적용한 Top5 표: {n}개")
print("saved:", DST)
