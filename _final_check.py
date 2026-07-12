# -*- coding: utf-8 -*-
"""최종 보고서(docx·PDF) 오류 종합 점검 — JSON 입력·필터·서식과 대조."""
import json, os, re, zipfile, glob, pickle
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from pypdf import PdfReader

HERE = os.path.dirname(os.path.abspath(__file__))
S = os.environ.get("COMPA_SCRATCH", "/private/tmp/claude-501/-Users-osung-work-compa/d6ed121c-12e4-45b4-b2fb-535b7554627c/scratchpad")
AP = os.environ.get("COMPA_APOLLO_DIR", "/Users/osung/work/apollo")
DOCX = "COMPA_필터전체_보고서.docx"; PDF = "COMPA_필터전체_보고서.pdf"
def norm(s): return re.sub(r"\s+", "", str(s or ""))
OK, BAD = [], []
def chk(cond, msg): (OK if cond else BAD).append(msg)

jb = json.load(open(os.path.join(HERE, "COMPA_통합best.json"), encoding="utf-8"))
pidf = json.load(open(f"{S}/pid_fields.json", encoding="utf-8"))
pats = json.load(open(f"{S}/pid_patents.json", encoding="utf-8"))
field6t = json.load(open(f"{S}/demand_field.json", encoding="utf-8"))

# ---- 0) JSON 정합성 ----
tops = [(k, t) for k, e in jb.items() for t in e["top5"]]
chk(len(jb) == 78, f"수요 78건 (실제 {len(jb)})")
chk(len(tops) == 390, f"추천 390건 (실제 {len(tops)})")
chk(all(len(e["top5"]) == 5 for e in jb.values()), "모든 수요 Top5=5")
empty_reason = [f"{k}:{t['과제명'][:15]}" for k, t in tops if not str(t.get("판단근거", "")).strip()]
chk(not empty_reason, f"빈 매칭근거 {len(empty_reason)}건")
empty_detail = [f"{k}" for k, t in tops if not str(t.get("추천근거_상세", "")).strip()]
chk(not empty_detail, f"빈 상세근거 {len(empty_detail)}건")
# 상세근거 4섹션 존재
SEC = ["연관성", "수요기술 사양 적합성", "추천 과제의 우수성", "유사 사례 및 실적"]
missing_sec = [k for k, t in tops if not all(f"[{s}]" in t.get("추천근거_상세", "") for s in SEC)]
chk(not missing_sec, f"4섹션 누락 {len(missing_sec)}건 {missing_sec[:5]}")
# 부정 톤 매칭근거(표)
NEG = re.compile(r"부재|미포함|불일치|미흡|아님|다름|없음|불가")
negs = [f"{k}:{t['판단근거'][:20]}" for k, t in tops if NEG.search(t.get("판단근거", ""))]
chk(not negs, f"부정 톤 매칭근거 {len(negs)}건 {negs[:3]}")

# ---- 1) 필터 준수(제출년도>=2020, 연구수행주체 4종) ----
ALLOW = {"대학", "출연연구소", "국공립연구소", "정부부처"}
subj_bad = [t["과제고유번호"] for k, t in tops if pidf.get(str(t["과제고유번호"]), {}).get("연구수행주체") not in ALLOW]
chk(not subj_bad, f"연구수행주체 위배 {len(subj_bad)}건")
print("· 제출년도 필터는 embeddings 로드해 확인 중…", flush=True)
try:
    emb = pd.read_pickle(f"{AP}/public_RnD_embeddings_pro_with_desc_260708.pkl")
    emb["과제고유번호"] = emb["과제고유번호"].astype(str)
    y = dict(zip(emb["과제고유번호"], pd.to_numeric(emb["제출년도"], errors="coerce")))
    del emb
    yr_bad = [t["과제고유번호"] for k, t in tops if not (y.get(str(t["과제고유번호"]), 0) >= 2020)]
    chk(not yr_bad, f"제출년도<2020 위배 {len(yr_bad)}건 {yr_bad[:5]}")
except Exception as e:
    BAD.append(f"제출년도 확인 실패: {e}")

# ---- 2) 특허 국가명(코드 잔존 금지) + 등록번호 ----
RAWCODE = re.compile(r"^(KR|US|CN|JP|EP|WO|XI|XU)$")
raw = [(pid, x["국가"]) for pid, v in pats.items() for x in v if RAWCODE.match(str(x["국가"]))]
chk(not raw, f"국가코드 미변환 {len(raw)}건 {raw[:5]}")
reg_no_num = [pid for pid, v in pats.items() for x in v if x["상태"] == "등록" and not x["등록번호"]]
chk(not reg_no_num, f"등록인데 등록번호 없음 {len(reg_no_num)}건")

# ---- 3) docx 구조/서식 ----
d = Document(DOCX)
try: zipfile.ZipFile(DOCX).testzip(); chk(True, "docx zip 무결성")
except Exception as e: BAD.append(f"docx 깨짐: {e}")
h2 = [p for p in d.paragraphs if p.style and p.style.name == "Heading 2" and "수요" in p.text]
h3 = [p for p in d.paragraphs if p.style and p.style.name == "Heading 3" and "TOP" in p.text.upper()]
chk(len(h2) == 78, f"docx 수요 제목 78 (실제 {len(h2)})")
chk(len(h3) == 390, f"docx TOP 제목 390 (실제 {len(h3)})")
# 볼드: 표지 타이틀 + H2 이름 + H3 이름
title_bold = any(r.bold and "매칭데이" in (r.text or "") for p in d.paragraphs for r in p.runs)
chk(title_bold, "표지 타이틀 볼드")
def name_bold(paras):  # 배지 아닌 이름 run 이 볼드인지
    bad = 0
    for p in paras:
        runs = [r for r in p.runs if r.text and not re.match(r"^\s*(수요|TOP)\s*\d*\s*$", r.text)]
        if runs and not any(r.bold for r in runs): bad += 1
    return bad
chk(name_bold(h2) == 0, f"수요기술명 볼드 아님 {name_bold(h2)}건")
chk(name_bold(h3) == 0, f"TOP 과제명 볼드 아님 {name_bold(h3)}건")
# 과제 정보표(4열) 390 + 특허 섹션 390
info_t = [t for t in d.tables if t.rows and t.rows[0].cells[0].text.strip() == "과제고유번호"]
chk(len(info_t) == 390, f"과제 정보표 390 (실제 {len(info_t)})")
chk(all(len(t.columns) == 4 for t in info_t), "정보표 모두 4열")
pat_hdr = sum(1 for t in d.tables if t.rows and [c.text.strip() for c in t.rows[0].cells][:2] == ["구분", "특허명"])
none_txt = sum(1 for p in d.paragraphs if p.text.strip() == "특허 실적 없음")
chk(pat_hdr + none_txt == 390, f"특허 실적(표{pat_hdr}+없음{none_txt})=390")
# 폴백/누출 텍스트 금지
alltext = "\n".join(p.text for p in d.paragraphs)
chk("과제 개요" not in alltext, "'과제 개요' 폴백 잔존")
chk("(출처" not in alltext, "'(출처' 잔존")
chk("<font" not in alltext, "'<font' 태그 누출")

# ---- 4) docx 데이터 대조: 수요별 정보표 pid 집합 == JSON pid 집합(동일제목 별개과제 허용) ----
from docx.text.paragraph import Paragraph; from docx.table import Table
body = list(d.element.body)
jb_by_dem = {}  # (기업,수요) -> [pid...]
for k, e in jb.items():
    jb_by_dem[(norm(e["기업명"]), norm(e["수요기술명"]))] = sorted(str(t["과제고유번호"]) for t in e["top5"])
doc_by_dem = {}; cur_comp = cur_dem = None
for ch in body:
    if ch.tag == qn("w:p"):
        p = Paragraph(ch, d); t = p.text.strip(); st = p.style.name if p.style else ""
        if st == "Heading 2" and "수요" in t: cur_dem = re.sub(r"^\s*수요\s*\d+\s*", "", t).strip(); cur_comp = None
    elif ch.tag == qn("w:tbl"):
        tb = Table(ch, d); h0 = tb.rows[0].cells[0].text.strip()
        if h0 == "기업명" and cur_comp is None: cur_comp = tb.rows[0].cells[1].text.strip()
        elif h0 == "과제고유번호":
            doc_by_dem.setdefault((norm(cur_comp), norm(cur_dem)), []).append(tb.rows[0].cells[1].text.strip())
mism = sum(1 for kk, v in jb_by_dem.items() if sorted(doc_by_dem.get(kk, [])) != v)
chk(mism == 0, f"docx 수요별 정보표 pid집합 불일치 {mism}건")

# ---- 5) PDF ----
try:
    r = PdfReader(PDF); npg = len(r.pages)
    chk(npg > 400, f"PDF 페이지 {npg}")
    txt0 = r.pages[0].extract_text()
    chk("COMPA 매칭데이" in txt0, "PDF 표지 제목")
    allpdf = "".join(p.extract_text() for p in r.pages[:20])
    chk("<font" not in allpdf, "PDF '<font' 누출")
except Exception as e:
    BAD.append(f"PDF 오류: {e}")

# ---- 결과 ----
print("\n===== 최종 보고서 점검 =====")
print(f"PASS {len(OK)} / FAIL {len(BAD)}")
for m in OK: print("  ✓", m)
if BAD:
    print("\n[FAIL]")
    for m in BAD: print("  ✗", m)
else:
    print("\n✅ 오류 없음")
