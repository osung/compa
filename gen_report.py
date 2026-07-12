# -*- coding: utf-8 -*-
"""COMPA 매칭데이 기술수요조사 최종 매칭 보고서(docx) 생성 — 출판물 품질 디자인.

입력:
  COMPA_통합best.json                                   최종 추천 Top5 + 근거(모델 생성)
  public_RnD_embeddings_pro_260601_with_desc_260708.pkl 연구개발단계/연구수행주체 (→ scratchpad/pid_fields.json 캐시)
  260625_..._분야별_78건.pdf                              78건 수요의 6T 분야 분류 (→ scratchpad/demand_field.json 캐시)
  COMPA_보고서_작성규칙.md                                 서식 규칙

디자인: 표제/표=Noto Sans KR, 본문 프로즈=Noto Serif KR. 네이비 계열 팔레트.
근거 텍스트(매칭 근거=판단근거, 상세 4섹션=추천근거_상세)는 JSON 모델 생성분 그대로 사용.
"""
import json, re, os, glob
from docx import Document
from docx.shared import Pt, RGBColor, Twips, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SCRATCH = os.environ.get("COMPA_SCRATCH", "/tmp/claude-1000/-mnt-d-work-compa/4bf38e1a-f59e-4115-8fee-5bed82365170/scratchpad")
OUT_BASE = "COMPA_매칭데이_기술수요조사_최종매칭_보고서"
PUBLISH_DATE = "2026. 7. 8."

def next_out_path():
    """매 실행마다 새 파일명(버전 자동 증가) — 기존 파일이 열려 있어도 충돌 없이 저장."""
    vers = [int(m.group(1)) for p in glob.glob(os.path.join(HERE, OUT_BASE + "_v*.docx"))
            if (m := re.search(r"_v(\d+)\.docx$", p))]
    n = (max(vers) + 1) if vers else 1
    return os.path.join(HERE, f"{OUT_BASE}_v{n}.docx")

OUT = os.environ.get("COMPA_REPORT_OUT", next_out_path())

# ---- 디자인 토큰 ------------------------------------------------------------
SANS = "Noto Sans KR"     # 표제·표·라벨 (사용자 Windows에 설치됨)
SERIF = "Noto Serif KR"   # 본문 프로즈 (개요·상세 매칭 근거)

INK      = "1B2430"   # 본문 먹색
NAVY     = "14315C"   # 주색 (표지·장·H1)
BLUE     = "2C5FA0"   # 보조 (H2·섹션 태그)
ACCENT   = "0E7C86"   # 강조 (수요 배지·캡션 포인트)
MUTED    = "6B7683"   # 캡션·부가 정보
HAIR     = "C9D2DE"   # 얇은 괘선
HEAD_BG  = "14315C"   # 표 헤더(진한 네이비, 흰 글자)
HEAD_FG  = "FFFFFF"
LABEL_BG = "EAF0F7"   # 라벨 셀 옅은 블루
ZEBRA    = "F5F8FC"   # 짝수행 옅은 배경
DISC_BD, DISC_BG, DISC_H, DISC_B = "C0392B", "FBEEED", "A93226", "7B241C"

# 과제 상세 1페이지 수렴 조판값(Serif 본문 기준으로 최장 블록까지 검증)
D_FONT = 8.6      # 상세 정보표/적합성/상세근거 본문
D_LS   = 1.1      # 상세근거 줄간격
SEC_AFTER = 2     # 섹션 간격(pt)
MARGIN_TB, MARGIN_LR = 0.7, 0.9  # inch

FIELD_TITLE = {"BT": "바이오기술 (BT) 분야", "IT": "정보기술 (IT) 분야",
               "NT": "나노기술 (NT) 분야", "ET": "환경기술 (ET) 분야", "융합": "융합기술 분야"}
FIELD_EN = {"BT": "BIOTECHNOLOGY", "IT": "INFORMATION TECHNOLOGY", "NT": "NANOTECHNOLOGY",
            "ET": "ENVIRONMENT TECHNOLOGY", "융합": "CONVERGENCE TECHNOLOGY"}
FIELD_ORDER = ["BT", "IT", "NT", "ET", "융합"]

DISCLAIMER = ("본 보고서는 APOLLO 인공지능을 이용해서 생성한 보고서로 사실과 다르거나 오류가 있을 수 있습니다. "
              "참고용으로만 활용하시고, 정확한 정보는 관련 자료를 통해 확인하시기 바랍니다. "
              "본 보고서는 AI 생성 내용의 정확성을 보증하지 않으며, 이를 근거로 한 판단 의사결정의 책임은 이용자에게 있습니다.")

# ---- 데이터 헬퍼 ------------------------------------------------------------
def load_data():
    demands = json.load(open(os.environ.get("COMPA_REPORT_JSON", os.path.join(HERE, "COMPA_통합best.json")), encoding="utf-8"))
    fields = json.load(open(os.path.join(SCRATCH, "demand_field.json"), encoding="utf-8"))
    pidf = json.load(open(os.path.join(SCRATCH, "pid_fields.json"), encoding="utf-8"))
    return demands, fields, pidf

def extract_period(desc):
    m = re.search(r'(\d{4})년\s*\d{1,2}월\s*\d{1,2}일에 시작.*?(\d{4})년\s*\d{1,2}월\s*\d{1,2}일에 종료', desc)
    if m:
        s, e = m.group(1), m.group(2)
        return (s,) if s == e else (s, e)
    m2 = re.search(r'(\d{4})년', desc)
    return (m2.group(1),) if m2 else None

def extract_class(desc):
    pats = [
        r'[,，]\s*([가-힣][가-힣·.\s]*?)\s*분야에\s*(?:속|해당)',
        r'수행(?:하며|했으며|하는 이 과제는|한 이 과제는)?\s*([가-힣][가-힣·.\s]{1,18}?)\s*분야에\s*(?:속|해당)',
        r'이 과제는\s*([가-힣][가-힣·.\s]{1,18}?)\s*분야에\s*(?:속|해당)',
        r'에서는?\s*([가-힣][가-힣·.\s]{1,18}?)\s*분야(?:의|에서)',
        r'[는은]\s*([가-힣][가-힣·.\s]{1,18}?)\s*분야(?:에서|의|에\s*속|에\s*해당)',
        r'([가-힣][가-힣·.\s]{1,18}?)\s*분야에\s*(?:속|해당)',
    ]
    for p in pats:
        m = re.search(p, desc)
        if m:
            v = re.sub(r'\s*\.\s*', '·', m.group(1).strip().rstrip('.').strip()).strip('·').strip()
            if 1 < len(v) <= 25 and not any(x in v for x in ("과제", "수행", "연구비", "총연구")):
                return v
    return None

def fmt_period(desc):
    p = extract_period(desc)
    return None if not p else (f"{p[0]}년" if len(p) == 1 else f"{p[0]}년 ~ {p[1]}년")

def year_lines(desc):
    p = extract_period(desc)
    if not p:
        return [""]
    return [p[0]] if len(p) == 1 else [f"{p[0]} ~", p[1]]

def split_sections(detail):
    secs, parts, i = [], re.split(r'(\[[^\]]+\])', detail), 1
    while i < len(parts):
        secs.append((parts[i].strip(), parts[i + 1].strip() if i + 1 < len(parts) else ""))
        i += 2
    return secs

# ---- 서식 저수준 ------------------------------------------------------------
def style_run(run, size=None, bold=False, color=None, family=SANS, spacing=None, italic=False):
    rPr = run._element.get_or_add_rPr()
    rf = rPr.get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(a), family)
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if spacing is not None:
        sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), str(spacing)); rPr.append(sp)

def run_shade(run, fill):
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill); rPr.append(shd)

def para(container, text="", size=10, bold=False, color=INK, align=None, family=SANS,
         before=None, after=None, line=None, indent=None, spacing=None, italic=False):
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    if before is not None: pf.space_before = Pt(before)
    if after is not None:  pf.space_after = Pt(after)
    if line is not None:   pf.line_spacing = line
    if indent is not None: pf.left_indent = Twips(indent)
    if text:
        style_run(p.add_run(text), size, bold, color, family, spacing, italic)
    return p

def para_border(p, edge="bottom", color=NAVY, sz=8, space=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr"); pPr.append(pbdr)
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), str(space)); e.set(qn("w:color"), color)
    pbdr.append(e)

def _tcPr(cell): return cell._tc.get_or_add_tcPr()

def set_cell_width(cell, dxa):
    tcPr = _tcPr(cell)
    for e in tcPr.findall(qn("w:tcW")): tcPr.remove(e)
    w = OxmlElement("w:tcW"); w.set(qn("w:w"), str(dxa)); w.set(qn("w:type"), "dxa"); tcPr.append(w)

def shade_cell(cell, fill):
    tcPr = _tcPr(cell)
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill); tcPr.append(shd)

def cell_vcenter(cell):
    va = OxmlElement("w:vAlign"); va.set(qn("w:val"), "center"); _tcPr(cell).append(va)

def table_grid(table, color=HAIR, sz=4, sides="all"):
    """sides: 'all'(전체 얇은 괘선) | 'h'(수평만) """
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    edges = ("top", "bottom", "insideH") if sides == "h" else \
            ("top", "left", "bottom", "right", "insideH", "insideV")
    for edge in edges:
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)

def table_cellmar(table, top=46, bottom=46, left=100, right=100):
    tblPr = table._tbl.tblPr
    cm = OxmlElement("w:tblCellMar")
    for side, v in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        e = OxmlElement(f"w:{side}"); e.set(qn("w:w"), str(v)); e.set(qn("w:type"), "dxa"); cm.append(e)
    tblPr.append(cm)

def table_fixed(table, widths):
    table.autofit = False; table.allow_autofit = False
    tblPr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblLayout"):          # 기존(autofit w=0) 요소 제거 후 재설정(중복 방지)
        for el in tblPr.findall(qn(tag)): tblPr.remove(el)
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed"); tblPr.append(layout)
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc in list(grid): grid.remove(gc)
        for w in widths:
            gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)
    for row in table.rows:
        for c, w in zip(row.cells, widths):
            set_cell_width(c, w)

def cell_indent(cell, left=200):
    for p in cell.paragraphs:
        p.paragraph_format.left_indent = Twips(left)

def rows_cantsplit(table):
    """각 행이 페이지 경계에서 쪼개지지 않도록(잔여 조각 방지)."""
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cs = OxmlElement("w:cantSplit"); cs.set(qn("w:val"), "true"); trPr.append(cs)

def fill_cell(cell, text, size=9, bold=False, align=None, color=INK, family=SANS, line=None):
    cell.text = ""
    p = cell.paragraphs[0]
    for i, ln in enumerate(str(text).split("\n")):
        if i > 0:
            p = cell.add_paragraph()
        if align is not None: p.alignment = align
        if line is not None: p.paragraph_format.line_spacing = line
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        style_run(p.add_run(ln), size, bold, color, family)

# ---- 헤더/푸터/페이지번호 ---------------------------------------------------
def add_page_field(paragraph, prefix="— ", suffix=" —"):
    if prefix:
        style_run(paragraph.add_run(prefix), 9, color=MUTED)
    r = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = " PAGE "
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    r._r.append(b); r._r.append(i); r._r.append(e)
    style_run(r, 9.5, bold=True, color=NAVY)
    if suffix:
        style_run(paragraph.add_run(suffix), 9, color=MUTED)

def set_pgnum_start(section, start=1):
    sectPr = section._sectPr
    e = sectPr.find(qn("w:pgNumType"))
    if e is None:
        e = OxmlElement("w:pgNumType"); sectPr.append(e)
    e.set(qn("w:start"), str(start))

# ---- 문서 골격 --------------------------------------------------------------
def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = SANS; normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    rf = normal.element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"): rf.set(qn(a), SANS)
    for hs in ("Heading 1", "Heading 2", "Heading 3"):
        st = doc.styles[hs]
        st.font.name = SANS
        st.element.get_or_add_rPr().get_or_add_rFonts()
        for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            st.element.rPr.rFonts.set(qn(a), SANS)

def set_margins(sec):
    sec.top_margin = Inches(MARGIN_TB); sec.bottom_margin = Inches(MARGIN_TB)
    sec.left_margin = Inches(MARGIN_LR); sec.right_margin = Inches(MARGIN_LR)
    sec.header_distance = Inches(0.4); sec.footer_distance = Inches(0.35)

def build():
    global LOGO_PATH
    demands, fields, pidf = load_data()
    try:
        LOGO_PATH = make_top_logo()
    except Exception as e:
        print("로고 크롭 실패(로고 없이 진행):", e); LOGO_PATH = None
    doc = Document()
    setup_styles(doc)
    set_margins(doc.sections[0])

    demand_by_field = {f: [] for f in FIELD_ORDER}
    for k in sorted(demands, key=int):
        demand_by_field[fields[k]].append(k)
    n_fields = len(set(fields.values()))
    total_recs = sum(len(v["top5"]) for v in demands.values())
    n_proj = len({t["과제고유번호"] for v in demands.values() for t in v["top5"]})

    build_cover(doc, len(demands), total_recs, n_proj)
    build_intro_toc(doc, demand_by_field, len(demands), total_recs, n_proj, n_fields)

    # ===== 본문 구역 =====
    doc.add_section(WD_SECTION.NEW_PAGE)
    body = doc.sections[-1]
    set_margins(body)
    set_pgnum_start(body, 1)
    setup_running_header(body)
    body.footer.is_linked_to_previous = False
    fp = body.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)
    doc.sections[0].footer.is_linked_to_previous = False
    setup_cover_header(doc.sections[0])   # 표지·목차에도 상단 로고(매 페이지)

    for i, f in enumerate([f for f in FIELD_ORDER if demand_by_field[f]]):
        build_chapter(doc, i + 1, f, demand_by_field[f], demands, pidf, first=(i == 0))

    doc.save(OUT)
    print("saved:", OUT)
    return OUT

LOGO_W = Inches(1.15)     # 헤더 로고 폭(높이 ≈ 0.245in)
LOGO_PATH = None          # build() 에서 크롭 후 설정

def make_top_logo():
    """'APOLLO 로고.png' 에서 위쪽(밝은 배경용) 로고만 잘라 저장."""
    import numpy as np
    from PIL import Image
    src = os.path.join(HERE, "APOLLO 로고.png")
    out = os.path.join(SCRATCH, "apollo_top.png")
    im = Image.open(src).convert("RGBA")
    alpha = np.array(im)[:, :, 3]
    rows = np.where((alpha > 16).sum(axis=1) > 5)[0]
    splits = np.where(np.diff(rows) > 20)[0]
    band = rows[:splits[0] + 1] if len(splits) else rows      # 위쪽 밴드
    y0, y1 = int(band[0]), int(band[-1])
    cols = np.where((alpha[y0:y1 + 1, :] > 16).sum(axis=0) > 2)[0]
    x0, x1 = int(cols[0]), int(cols[-1])
    pad = 12
    im.crop((max(0, x0 - pad), max(0, y0 - pad),
             min(im.width, x1 + pad), min(im.height, y1 + pad))).save(out)
    return out

def _hdr_bottom_rule(tbl):
    borders = OxmlElement("w:tblBorders")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6"); b.set(qn("w:space"), "0"); b.set(qn("w:color"), HAIR)
    borders.append(b); tbl._tbl.tblPr.append(borders)

def setup_running_header(section):
    """본문 구역: 좌측 러닝 타이틀 + 우측(맨 오른쪽) APOLLO 로고 + 하단 얇은 괘선.
    탭 정렬은 렌더러(LibreOffice)에서 이미지가 가운데로 밀리는 문제가 있어 2열 표로 확정 배치."""
    section.header.is_linked_to_previous = False
    hdr = section.header
    for p in list(hdr.paragraphs):        # 기본 빈 문단 제거(헤더 높이 최소화)
        p._element.getparent().remove(p._element)
    tbl = hdr.add_table(rows=1, cols=2, width=Inches(6.7))
    tbl.autofit = False
    left, right = tbl.rows[0].cells
    lp = left.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_before = Pt(0); lp.paragraph_format.space_after = Pt(0)
    style_run(lp.add_run("COMPA  기술수요–공공 R&D 매칭 보고서"), 8, color=MUTED, spacing=20)
    cell_vcenter(left)
    rp = right.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT   # 로고를 셀 오른쪽 끝으로
    rp.paragraph_format.space_before = Pt(0); rp.paragraph_format.space_after = Pt(0)
    if LOGO_PATH:
        rp.add_run().add_picture(LOGO_PATH, width=LOGO_W)
    cell_vcenter(right)
    table_fixed(tbl, [5900, 3748]); table_cellmar(tbl, 0, 0, 0, 0)
    _hdr_bottom_rule(tbl)
    tp = hdr.add_paragraph()               # 표 뒤 최소 문단(높이 무시 가능)
    tp.paragraph_format.space_before = Pt(0); tp.paragraph_format.space_after = Pt(0)
    tp.paragraph_format.line_spacing = Pt(2)
    style_run(tp.add_run(""), 1)

def setup_cover_header(section):
    """표지·목차 구역: 우측 상단 로고만(러닝 타이틀 없음)."""
    section.header.is_linked_to_previous = False
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if LOGO_PATH:
        hp.add_run().add_picture(LOGO_PATH, width=LOGO_W)

# ---- 표지 -------------------------------------------------------------------
def build_cover(doc, n_dem, n_rec, n_proj):
    para(doc, "", after=54)
    para(doc, "TECHNOLOGY  DEMAND  ×  PUBLIC  R&D  MATCHING", 10.5, bold=True,
         color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=60, after=10)
    para(doc, "COMPA 매칭데이", 27, bold=True, color=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "기술수요조사 최종 매칭 보고서", 27, bold=True, color=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    rule = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    para_border(rule, "bottom", NAVY, 18, 2)
    para(doc, "기업 진성수요 × 공공 R&D 과제, 의미 기반 매칭 결과",
         12, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=44)

    # 메타 카드(무테 2열 표)
    t = doc.add_table(rows=2, cols=2)
    t.alignment = 1  # center
    stats = [("대상 수요기술", f"{n_dem}건"), ("추천 과제", f"{n_rec}건")]
    for j, (lab, val) in enumerate(stats):
        fill_cell(t.rows[0].cells[j], val, 20, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(t.rows[1].cells[j], lab, 9.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    table_fixed(t, [3200, 3200]); table_cellmar(t, 30, 30, 60, 60)
    para(doc, "", after=40)
    para(doc, f"발행일  {PUBLISH_DATE}      생성  APOLLO AI 매칭 엔진",
         10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)

    para(doc, "", after=30)
    add_disclaimer_box(doc)

def add_disclaimer_box(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "6")
        e.set(qn("w:space"), "10"); e.set(qn("w:color"), DISC_BD)
        pbdr.append(e)
    pPr.append(pbdr)
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), DISC_BG); pPr.append(shd)
    r1 = p.add_run("※  유의사항"); style_run(r1, 10.5, bold=True, color=DISC_H); r1.add_break()
    style_run(p.add_run(DISCLAIMER), 9.5, color=DISC_B, family=SANS)

# ---- 개요 · 목차 ------------------------------------------------------------
def section_label(doc, text, before=14, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    bar = p.add_run("▍"); style_run(bar, 13, bold=True, color=ACCENT)
    style_run(p.add_run(" " + text), 13.5, bold=True, color=NAVY)
    return p

def build_intro_toc(doc, dbf, n_dem, n_rec, n_proj, n_fields):
    doc.add_paragraph().paragraph_format.page_break_before = True
    section_label(doc, "개요", before=4)
    overview = (
        f"본 보고서는 매칭데이 기술수요조사를 통해 취합된 기업 진성수요 {n_dem}건을 대상으로, "
        f"공공 R&D 과제 데이터베이스와의 의미 기반 매칭을 수행한 결과를 정리한 것이다. "
        f"각 수요기술에 대해 적합도가 높은 추천 과제 상위 5건(총 {n_rec}건, 중복 제외 {n_proj}개 과제)을 "
        f"선정하고, 매칭 근거와 상세 추천 근거를 함께 제시하였다. "
        f"수요는 6T 기술 분류에 따라 {n_fields}개 분야(바이오기술·정보기술·나노기술·환경기술·융합기술)로 구분하여 수록하였다."
    )
    para(doc, overview, 10.5, color=INK, family=SERIF, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.6, after=8)
    para(doc, "각 수요기술은 다음 순서로 구성된다.", 10.5, color=INK, family=SERIF, after=3)
    for ln in ["수요 정보  —  기업명 · 수요기술 내용 · 수요기술 사양",
               "최종 추천 과제 Top 5  —  순위 · 과제명 · 수행기관 · 과제수행년도 · 매칭 근거",
               "추천 과제별 상세 정보표 및 상세 매칭 근거  —  연관성 · 수요기술 사양 적합성 · 추천 과제의 우수성 · 유사 사례 및 실적"]:
        q = doc.add_paragraph(); q.paragraph_format.left_indent = Twips(260); q.paragraph_format.space_after = Pt(3)
        style_run(q.add_run("· "), 10.5, bold=True, color=ACCENT)
        style_run(q.add_run(ln), 10, color=INK, family=SERIF)

    section_label(doc, "목차", before=20)
    W = 9550  # 우측 탭 위치(dxa)
    for i, f in enumerate([f for f in FIELD_ORDER if dbf[f]], 1):
        ks = dbf[f]
        rng = f"수요 {ks[0]}–{ks[-1]}" if len(ks) > 1 else f"수요 {ks[0]}"
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.tab_stops.add_tab_stop(Twips(W), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        style_run(p.add_run(f"제 {i} 장"), 10.5, bold=True, color=ACCENT)
        style_run(p.add_run(f"   {FIELD_TITLE[f]}"), 11, bold=True, color=NAVY)
        style_run(p.add_run(f"\t{rng} · {len(ks)}건"), 9.5, color=MUTED)

# ---- 장(章) 오프너 ----------------------------------------------------------
def build_chapter(doc, no, f, ks, demands, pidf, first=False):
    h1 = doc.add_heading(level=1)
    if not first:
        h1.paragraph_format.page_break_before = True
    h1.paragraph_format.space_before = Pt(0); h1.paragraph_format.space_after = Pt(2)
    style_run(h1.add_run(f"제{no}장  {FIELD_TITLE[f]}"), 21, bold=True, color=NAVY)
    cap = para(doc, FIELD_EN[f], 9, bold=True, color=ACCENT, spacing=50, after=6)
    para_border(cap, "bottom", NAVY, 14, 6)
    para(doc, f"수요기술 {len(ks)}건  ·  수요 {ks[0]}–{ks[-1]}" if len(ks) > 1
         else f"수요기술 {len(ks)}건  ·  수요 {ks[0]}", 10, color=MUTED, after=12)

    # 수요기술 목록
    section_label(doc, "수요기술 목록", before=2, after=6)
    t = doc.add_table(rows=1, cols=3)
    table_grid(t, HAIR, 4, "all"); table_cellmar(t)
    for c, txt in zip(t.rows[0].cells, ("번호", "수요기술명", "기업명")):
        shade_cell(c, HEAD_BG); cell_vcenter(c)
        fill_cell(c, txt, 9.5, bold=True, color=HEAD_FG, align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx, k in enumerate(ks):
        row = t.add_row().cells
        if idx % 2 == 1:
            for c in row: shade_cell(c, ZEBRA)
        fill_cell(row[0], k, 9.5, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(row[1], demands[k]["수요기술명"], 9.5, line=1.25)
        fill_cell(row[2], demands[k]["기업명"], 9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        for c in row: cell_vcenter(c)
    table_fixed(t, [760, 5080, 2800])

    for k in ks:
        build_demand(doc, k, demands[k], pidf)

# ---- 수요 블록 --------------------------------------------------------------
def build_demand(doc, k, dm, pidf):
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.page_break_before = True
    h2.paragraph_format.space_before = Pt(0); h2.paragraph_format.space_after = Pt(5)
    # 텍스트는 "[수요 N] 제목" 유지, 접두 배지만 색/음영
    badge = h2.add_run(f" 수요 {k} ")
    style_run(badge, 12, bold=True, color="FFFFFF"); run_shade(badge, ACCENT)
    style_run(h2.add_run("  "), 12)
    style_run(h2.add_run(dm["수요기술명"]), 14, bold=True, color=NAVY)
    para_border(h2, "bottom", HAIR, 6, 6)

    # 수요 정보 표
    rows = [("기업명", dm.get("기업명", ""))]
    if (dm.get("수요기술 내용") or "").strip():
        rows.append(("수요기술 내용", dm["수요기술 내용"].strip()))
    if (dm.get("수요기술 사양") or "").strip():
        rows.append(("수요기술 사양", dm["수요기술 사양"].strip()))
    t = doc.add_table(rows=0, cols=2)
    table_grid(t, HAIR, 4, "all"); table_cellmar(t, 44, 44, 110, 110)
    for label, val in rows:
        cells = t.add_row().cells
        shade_cell(cells[0], LABEL_BG); cell_vcenter(cells[0])
        fill_cell(cells[0], label, 9.5, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
        align = WD_ALIGN_PARAGRAPH.JUSTIFY if label != "기업명" else None
        fam = SERIF if label != "기업명" else SANS
        fill_cell(cells[1], val, 9, align=align, family=fam, line=1.28)
    table_fixed(t, [1560, 7080])

    # 최종 추천 Top5
    section_label(doc, "최종 추천 과제  Top 5", before=9, after=5)
    t = doc.add_table(rows=1, cols=5)
    table_grid(t, HAIR, 4, "all"); table_cellmar(t, 32, 32, 90, 90)
    for c, txt in zip(t.rows[0].cells, ("순위", "과제명", "수행기관", "수행년도", "매칭 근거")):
        shade_cell(c, HEAD_BG); cell_vcenter(c)
        fill_cell(c, txt, 9, bold=True, color=HEAD_FG, align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx, tp in enumerate(dm["top5"]):
        cells = t.add_row().cells
        if idx % 2 == 1:
            for c in cells: shade_cell(c, ZEBRA)
        fill_cell(cells[0], str(tp["rank"]), 11, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(cells[1], tp["과제명"], 8.8, line=1.18)
        fill_cell(cells[2], tp.get("수행기관", ""), 8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(cells[3], "\n".join(year_lines(tp.get("과제설명문", ""))), 8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(cells[4], tp.get("판단근거", ""), 8.8, align=WD_ALIGN_PARAGRAPH.JUSTIFY, family=SERIF, line=1.18)
        for c in cells: cell_vcenter(c)
    table_fixed(t, [640, 3360, 1500, 1040, 2100])
    rows_cantsplit(t)

    for tp in dm["top5"]:
        build_top_detail(doc, tp, pidf)

# ---- 과제 상세 블록 (1페이지) -----------------------------------------------
def build_top_detail(doc, tp, pidf):
    pid = str(tp["과제고유번호"])
    h3 = doc.add_heading(level=3)
    h3.paragraph_format.page_break_before = True
    h3.paragraph_format.space_before = Pt(0); h3.paragraph_format.space_after = Pt(6)
    pill = h3.add_run(f" TOP {tp['rank']} ")
    style_run(pill, 11, bold=True, color="FFFFFF"); run_shade(pill, NAVY)
    style_run(h3.add_run("  "), 12)
    style_run(h3.add_run(tp["과제명"]), 12.5, bold=True, color=INK)
    para_border(h3, "bottom", HAIR, 6, 6)

    extra = pidf.get(pid, {})
    info = [("과제고유번호", pid)]
    if (per := fmt_period(tp.get("과제설명문", ""))): info.append(("과제수행기간", per))
    if (cls := extract_class(tp.get("과제설명문", ""))): info.append(("과학기술표준분류(중)", cls))
    if extra.get("연구개발단계"): info.append(("연구개발단계", extra["연구개발단계"]))
    info.append(("과제수행기관", tp.get("수행기관", "")))
    if extra.get("연구책임자명"): info.append(("연구책임자", extra["연구책임자명"]))
    if extra.get("국가연구자번호"): info.append(("국가연구자번호", extra["국가연구자번호"]))
    if extra.get("연구수행주체"): info.append(("연구수행주체", extra["연구수행주체"]))

    # 4열(라벨|값|라벨|값)로 묶어 행 수를 절반으로 축소
    t = doc.add_table(rows=0, cols=4)
    table_grid(t, HAIR, 4, "all"); table_cellmar(t, 34, 34, 110, 110)
    for i in range(0, len(info), 2):
        pairs = [info[i]] + ([info[i + 1]] if i + 1 < len(info) else [("", "")])
        cells = t.add_row().cells
        for j, (label, val) in enumerate(pairs):
            lc, vc = cells[j * 2], cells[j * 2 + 1]
            cell_vcenter(lc); cell_vcenter(vc)
            if label:
                shade_cell(lc, LABEL_BG)
                fill_cell(lc, label, D_FONT, bold=True, color=NAVY); cell_indent(lc, 80)
            fill_cell(vc, str(val), D_FONT)
    table_fixed(t, [2300, 2020, 2300, 2020])  # 라벨열=최장 라벨'과학기술표준분류(중)' 한 줄 최소폭

    # 적합성 판단 (강조 라인)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    tag = p.add_run(" 적합성 판단 "); style_run(tag, D_FONT, bold=True, color="FFFFFF"); run_shade(tag, BLUE)
    style_run(p.add_run("  " + tp.get("판단근거", "")), D_FONT + 0.5, bold=True, color=INK)

    para(doc, "상세 매칭 근거", D_FONT + 1, bold=True, color=NAVY, before=4, after=3)
    for title, body in split_sections(tp.get("추천근거_상세", "")):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(SEC_AFTER); p.paragraph_format.line_spacing = D_LS
        style_run(p.add_run(title.strip("[]") + "  "), D_FONT, bold=True, color=BLUE)
        style_run(p.add_run(body), D_FONT, color=INK, family=SERIF)

if __name__ == "__main__":
    build()
