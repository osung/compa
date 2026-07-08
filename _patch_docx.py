# -*- coding: utf-8 -*-
"""COMPA_최종Top5_보고서.docx 패치:
 (1) 표 헤더 '매칭 이유' → '매칭 근거'
 (2) 72건(상세근거 없던 Top)의 표 '매칭 근거' 셀 = 실제 매칭 근거 한 문장
 (3) 해당 Top 상세 블록: '과제 개요' → '적합성 판단' + '상세 매칭 근거' 4섹션
서식은 정상 항목의 문단을 복제해 유지."""
import json, re, copy, sys
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

SCRATCH = "/private/tmp/claude-501/-Users-osung-work-compa/d6ed121c-12e4-45b4-b2fb-535b7554627c/scratchpad"
SRC = sys.argv[1] if len(sys.argv) > 1 else "COMPA_최종Top5_보고서.docx"
DST = sys.argv[2] if len(sys.argv) > 2 else SRC

def norm(s): return re.sub(r"\s+", "", str(s)).strip()

SECN = ["연관성", "수요기술 사양 적합성", "추천 과제의 우수성", "유사 사례 및 실적"]

def split_sections(detail):
    parts = re.split(r"(?=\[(?:%s)\])" % "|".join(map(re.escape, SECN)), detail.strip())
    secs = []
    for p in parts:
        m = re.match(r"\[([^\]]+)\]\s*(.*)", p.strip(), re.DOTALL)
        if m and m.group(1).strip() in SECN:
            secs.append((m.group(1).strip(), re.sub(r"\s+", " ", m.group(2).strip())))
    return secs

# ---------- 타깃 맵 ----------
tgt = {}
data = json.load(open(f"{SCRATCH}/regen_targets.json"))
for r in data["recover"]:
    tgt[(norm(r["기업명"]), norm(r["과제명"]))] = {"reason": r["판단근거"].strip(),
                                                 "secs": split_sections(r["추천근거_상세"])}
regen_out = json.load(open(f"{SCRATCH}/regen_out.json"))
for k, v in regen_out.items():
    reason = (v.get("매칭근거_short") or v.get("판단근거") or "").strip()
    tgt[(norm(v["기업명"]), norm(v["과제명"]))] = {"reason": reason,
                                                 "secs": split_sections(v["추천근거_상세"])}
print("targets:", len(tgt))

d = Document(SRC)

# ---------- 템플릿 문단 확보(정상 항목에서 복제) ----------
tpl_fit = tpl_head = tpl_sec = None
for p in d.paragraphs:
    t = p.text.strip()
    if tpl_fit is None and t.startswith("적합성 판단"):
        tpl_fit = copy.deepcopy(p._p)
    if tpl_head is None and t.startswith("상세 매칭 근거"):
        tpl_head = copy.deepcopy(p._p)
    if tpl_sec is None and t.startswith("[연관성]"):
        tpl_sec = copy.deepcopy(p._p)
    if tpl_fit is not None and tpl_head is not None and tpl_sec is not None:
        break
assert tpl_fit is not None and tpl_head is not None and tpl_sec is not None, "template missing"

def clone_set(tpl, label=None, body=None):
    el = copy.deepcopy(tpl)
    para = Paragraph(el, None)
    runs = para.runs
    if label is not None and runs:
        runs[0].text = label
    if body is not None and runs:
        if len(runs) >= 2:
            runs[1].text = body
            for r in runs[2:]:
                r.text = ""
        else:
            runs[-1].text = body
    return el

def set_cell(cell, value):
    para = cell.paragraphs[0]
    runs = para.runs
    if not runs:
        para.add_run(value); return
    tgt_run = max(runs, key=lambda r: len(r.text)) if any(r.text for r in runs) else runs[-1]
    for r in runs:
        r.text = value if r is tgt_run else ""

# ---------- 본문 순서 순회 ----------
cur_comp = cur_top = None
overview_jobs = []   # (para_elem, key)
n_hdr = n_cell = n_block = 0

for child in list(d.element.body):
    if child.tag == qn("w:p"):
        para = Paragraph(child, d)
        t = para.text.strip()
        st = para.style.name if para.style is not None else ""
        if st == "Normal" and t.startswith("기업명:"):
            cur_comp = t.split("기업명:", 1)[1].strip()
        elif st == "Heading 3" and re.match(r"Top\d", t):
            m = re.match(r"Top\d+\.\s*(.*)", t)
            cur_top = re.sub(r"\s*\(출처[:：].*?\)\s*$", "", m.group(1)).strip()
        elif st == "Normal" and t.startswith("과제 개요"):
            key = (norm(cur_comp), norm(cur_top))
            if key in tgt:
                overview_jobs.append((child, key))
    elif child.tag == qn("w:tbl"):
        from docx.table import Table
        tb = Table(child, d)
        rows = tb.rows
        if not rows or len(rows[0].cells) < 4:
            continue
        hdr = [c.text.strip() for c in rows[0].cells]
        if hdr[0] == "순위" and hdr[3] in ("매칭 이유", "매칭 근거"):
            if hdr[3] == "매칭 이유":
                set_cell(rows[0].cells[3], "매칭 근거"); n_hdr += 1
            for r in rows[1:]:
                cells = r.cells
                jobname = cells[1].text.strip()
                key = (norm(cur_comp), norm(jobname))
                if key in tgt:
                    set_cell(cells[3], tgt[key]["reason"]); n_cell += 1

# ---------- 상세 블록 교체 ----------
for para_elem, key in overview_jobs:
    secs = tgt[key]["secs"]
    reason = tgt[key]["reason"]
    new_elems = [clone_set(tpl_fit, body=reason), clone_set(tpl_head)]
    for name, bodytext in secs:
        new_elems.append(clone_set(tpl_sec, label=f"[{name}] ", body=bodytext))
    for el in new_elems:
        para_elem.addprevious(el)
    para_elem.getparent().remove(para_elem)
    n_block += 1

d.save(DST)
print(f"헤더 개명:{n_hdr}  표셀 채움:{n_cell}  상세블록 교체:{n_block}")
print("saved:", DST)
