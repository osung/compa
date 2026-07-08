# -*- coding: utf-8 -*-
"""첫 페이지 제목 아래에 AI 생성 보고서 경고문 박스 삽입."""
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

SRC = sys.argv[1] if len(sys.argv) > 1 else "COMPA_최종Top5_보고서.docx"
DST = sys.argv[2] if len(sys.argv) > 2 else SRC

LEAD = "※ 유의사항"
BODY = ("본 보고서는 APOLLO 인공지능을 이용해서 생성한 보고서로 사실과 다르거나 오류가 있을 수 있습니다. "
        "참고용으로만 활용하시고, 정확한 정보는 관련 자료를 통해 확인하시기 바랍니다. "
        "본 보고서는 AI 생성 내용의 정확성을 보증하지 않으며, 이를 근거로 한 판단 의사결정의 책임은 이용자에게 있습니다.")

d = Document(SRC)

# 이미 삽입돼 있으면 중복 방지
if any("APOLLO 인공지능을 이용해서 생성" in p.text for p in d.paragraphs):
    print("이미 경고문이 있어 건너뜀"); d.save(DST); sys.exit(0)

title = d.paragraphs[0]
p = OxmlElement("w:p")

# 문단 속성: 테두리 + 음영 + 여백
pPr = OxmlElement("w:pPr")
pBdr = OxmlElement("w:pBdr")
for side in ("top", "left", "bottom", "right"):
    b = OxmlElement(f"w:{side}")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "12")
    b.set(qn("w:space"), "6"); b.set(qn("w:color"), "C00000")
    pBdr.append(b)
pPr.append(pBdr)
shd = OxmlElement("w:shd")
shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "FCE9E9")
pPr.append(shd)
spc = OxmlElement("w:spacing")
spc.set(qn("w:before"), "120"); spc.set(qn("w:after"), "120")
pPr.append(spc)
p.append(pPr)

def add_run(text, bold, size, color, break_before=False):
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if bold:
        rPr.append(OxmlElement("w:b"))
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rPr.append(sz)
    col = OxmlElement("w:color"); col.set(qn("w:val"), color); rPr.append(col)
    r.append(rPr)
    if break_before:
        r.append(OxmlElement("w:br"))
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    r.append(t)
    p.append(r)

add_run(LEAD, True, 10.5, "C00000")
add_run(BODY, False, 9.5, "7F1D1D", break_before=True)

title._p.addnext(p)
d.save(DST)
print("경고문 삽입 완료 →", DST)
