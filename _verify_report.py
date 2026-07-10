# -*- coding: utf-8 -*-
"""생성된 보고서 docx(출판물 디자인)를 원본 JSON 및 규칙(§1~§8)과 대조 검증."""
import json, re, zipfile, os, glob
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
_c = sorted(glob.glob(os.path.join(HERE, "COMPA_매칭데이_기술수요조사_최종매칭_보고서_v*.docx")),
            key=lambda p: int(re.search(r"_v(\d+)\.docx$", p).group(1)))
DOCX = _c[-1]
SCRATCH = "/tmp/claude-1000/-mnt-d-work-compa/4bf38e1a-f59e-4115-8fee-5bed82365170/scratchpad"
print("검증 대상:", os.path.basename(DOCX))

demands = json.load(open(os.path.join(HERE, "COMPA_통합best.json"), encoding="utf-8"))
fields = json.load(open(os.path.join(SCRATCH, "demand_field.json"), encoding="utf-8"))

PASS, FAIL = [], []
def chk(cond, msg): (PASS if cond else FAIL).append(msg)
def norm(s): return re.sub(r"\s+", "", s or "")

# 0. zip 무결성
chk(zipfile.ZipFile(DOCX).testzip() is None, "zip 무결성")

d = Document(DOCX)
paras, tables = d.paragraphs, d.tables

# 1. 계층 구조
h1 = [p for p in paras if p.style.name == "Heading 1"]
h2 = [p for p in paras if p.style.name == "Heading 2"]
h3 = [p for p in paras if p.style.name == "Heading 3"]
chk(len(h1) == 5, f"Heading1(분야)={len(h1)} (기대 5)")
chk(len(h2) == 78, f"Heading2(수요)={len(h2)} (기대 78)")
chk(len(h3) == 390, f"Heading3(TopN)={len(h3)} (기대 390)")
chk(len(tables) >= 551, f"표 개수={len(tables)} (본문 551 + 표지 메타 1)")
chk(len(d.sections) == 2, f"구역 수={len(d.sections)}")

# 2. 수요 제목 78건: 배지(수요 N) + 수요기술명 순서 일치
h2t = [norm(p.text) for p in h2]
exp = [norm(f"수요 {k} {demands[k]['수요기술명']}") for k in sorted(demands, key=int)]
mis = sum(a != b for a, b in zip(h2t, exp))
chk(mis == 0, f"수요 제목/순서 78건 일치 (불일치 {mis})")

# 3. 장 제목/순서 (공백 무시 비교)
FT = {"BT": "바이오기술 (BT) 분야", "IT": "정보기술 (IT) 분야", "NT": "나노기술 (NT) 분야",
      "ET": "환경기술 (ET) 분야", "융합": "융합기술 분야"}
exp_ch = [norm(f"제{i}장  {FT[f]}") for i, f in enumerate(["BT","IT","NT","ET","융합"], 1)]
chk([norm(p.text) for p in h1] == exp_ch, f"장 제목/순서: {[p.text for p in h1]}")

# 4. Top5 표: 78개, 데이터 JSON 일치
top5 = [t for t in tables if t.rows and t.rows[0].cells[0].text == "순위"]
chk(len(top5) == 78, f"Top5 표={len(top5)} (기대 78)")
bad = 0
for k, t in zip(sorted(demands, key=int), top5):
    body = t.rows[1:]
    if len(body) != 5: bad += 100; continue
    for row, tp in zip(body, demands[k]["top5"]):
        c = [x.text for x in row.cells]
        if c[0] != str(tp["rank"]) or norm(c[1]) != norm(tp["과제명"]) \
           or norm(c[2]) != norm(tp["수행기관"]) or norm(c[4]) != norm(tp["판단근거"]):
            bad += 1
chk(bad == 0, f"Top5 셀 데이터 JSON 일치 (불일치 {bad}/390)")

# 5. 과제 정보표 390개, 과제고유번호 순서/값 일치
info = [t for t in tables if t.rows and t.rows[0].cells[0].text == "과제고유번호"]
chk(len(info) == 390, f"과제 정보표={len(info)} (기대 390)")
jpids = [str(tp["과제고유번호"]) for k in sorted(demands, key=int) for tp in demands[k]["top5"]]
dpids = [t.rows[0].cells[1].text for t in info]
chk(dpids == jpids, f"과제고유번호 순서/값 일치 (불일치 {sum(a!=b for a,b in zip(dpids,jpids))})")

# 6. 상세 매칭 근거 구조: '상세 매칭 근거' 헤더 390, '적합성 판단' 배지 390, 섹션 4종×390
n_hdr = sum(1 for p in paras if p.text.strip() == "상세 매칭 근거")
n_fit = sum(1 for p in paras if any(r.text.strip() == "적합성 판단" for r in p.runs))
chk(n_hdr == 390, f"'상세 매칭 근거' 헤더={n_hdr} (기대 390)")
chk(n_fit == 390, f"'적합성 판단' 태그={n_fit} (기대 390)")
SECS = ["연관성", "수요기술 사양 적합성", "추천 과제의 우수성", "유사 사례 및 실적"]
sec_cnt = {s: 0 for s in SECS}
for p in paras:
    if p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY and p.runs:
        head = p.runs[0].text.strip()
        if head in sec_cnt:
            sec_cnt[head] += 1
for s in SECS:
    chk(sec_cnt[s] == 390, f"섹션 '{s}' 문단={sec_cnt[s]} (기대 390)")

# 7. 표 폭 규칙
def gw(t):
    return [int(gc.get(qn("w:w"))) for gc in t._tbl.find(qn("w:tblGrid"))]
def first(h0):
    return next(t for t in tables if t.rows and t.rows[0].cells[0].text == h0)
chk(gw(first("번호")) == [760, 5080, 2800], f"수요목록표 폭 {gw(first('번호'))}")
chk(gw(first("기업명")) == [1560, 7080], f"수요정보표 폭 {gw(first('기업명'))}")
chk(gw(top5[0]) == [640, 3360, 1500, 1040, 2100], f"Top5표 폭 {gw(top5[0])}")
chk(gw(info[0]) == [1900, 4300], f"과제정보표 폭 {gw(info[0])}")
for nm, t in [("수요목록", first("번호")), ("Top5", top5[0]), ("과제정보", info[0])]:
    chk(sum(gw(t)) <= 9648, f"{nm}표 폭합 {sum(gw(t))} ≤ 페이지폭")

# 8. 디자인 서식: 표 헤더 네이비/흰글자, 얇은 괘선, 고정 레이아웃
xml = top5[0]._tbl.xml
chk('w:fill="14315C"' in xml, "Top5 헤더 네이비(14315C) 음영")
chk("C9D2DE" in xml, "Top5 얇은 괘선(C9D2DE)")
chk('w:type="fixed"' in xml, "Top5 고정 레이아웃")
chk("w:tblCellMar" in xml, "Top5 셀 여백(padding)")
chk("w:cantSplit" in xml, "Top5 행 분할 방지(cantSplit)")
# 폰트: Noto Sans/Serif KR 사용
allxml = "".join(p._p.xml for p in paras[:400])
chk("Noto Sans KR" in allxml, "본문 폰트 Noto Sans KR 적용")
detail_xml = "".join(p._p.xml for p in paras if p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY)
chk("Noto Serif KR" in detail_xml, "상세근거 프로즈 Noto Serif KR 적용")

# 9. 양쪽맞춤
J = WD_ALIGN_PARAGRAPH.JUSTIFY
chk(top5[0].rows[1].cells[4].paragraphs[0].alignment == J, "Top5 '매칭 근거' 셀 justify")
djust = [p for p in paras if p.alignment == J and p.runs and p.runs[0].text.strip() in SECS]
chk(len(djust) == 1560, f"상세근거 섹션 문단 justify={len(djust)} (기대 1560)")

# 10. 페이지번호/헤더/푸터/유의사항
s2 = d.sections[1]
pg = s2._sectPr.find(qn("w:pgNumType"))
chk(pg is not None and pg.get(qn("w:start")) == "1", "구역2 페이지번호 start=1")
chk("PAGE" in s2.footer.paragraphs[0]._p.xml, "구역2 footer PAGE 필드")
chk("COMPA" in s2.header._element.xml, "구역2 러닝 헤더(표 안 타이틀)")
chk(d.sections[0].footer.paragraphs[0].text == "", "구역1(표지/목차) footer 비움")
disc = next((p for p in paras if "APOLLO 인공지능" in p.text and "유의사항" in p.text), None)
chk(disc is not None, "유의사항 문구 존재")
if disc:
    chk("pBdr" in disc._p.xml and "C0392B" in disc._p.xml, "유의사항 박스 테두리(C0392B)")

# 11. 표기 금지
full = "\n".join(p.text for p in paras)
chk("출처" not in full, "본문에 '출처' 표기 없음")

# 11-1. 표지 '중복 제외 과제' 카드 삭제
cell_texts = [c.text.strip() for t in tables for r in t.rows for c in r.cells]
chk("중복 제외 과제" not in cell_texts, "표지 '중복 제외 과제' 카드 삭제됨")

# 11-2. APOLLO 로고: 두 구역 헤더 모두 이미지 포함(→ 매 페이지 노출)
def header_has_image(sec):
    x = sec.header._element.xml           # 문단·표 포함 전체 헤더 XML
    return "blip" in x or "graphicData" in x
chk(header_has_image(d.sections[0]), "표지·목차 헤더에 로고 이미지(우측)")
chk(header_has_image(d.sections[1]), "본문 헤더에 로고 이미지(우측)")
media = [n for n in zipfile.ZipFile(DOCX).namelist() if n.startswith("word/media/")]
chk(len(media) >= 1, f"임베드 로고 미디어 파일={media}")

# 12. 페이지 나누기 속성
chk(all(p.paragraph_format.page_break_before for p in h2), "모든 수요(H2) page_break_before")
chk(all(p.paragraph_format.page_break_before for p in h3), "모든 과제(H3) page_break_before")

# 13. 수요 정보표 기업명 일치
dtabs = [t for t in tables if t.rows and t.rows[0].cells[0].text == "기업명"]
chk(len(dtabs) == 78, f"수요 정보표={len(dtabs)}")
ok = sum(1 for k, t in zip(sorted(demands, key=int), dtabs)
         if norm(t.rows[0].cells[1].text) == norm(demands[k]["기업명"]))
chk(ok == 78, f"수요 정보표 기업명 일치 {ok}/78")

print(f"\n{'='*58}\n검증 결과: PASS {len(PASS)} / FAIL {len(FAIL)}\n{'='*58}")
for m in PASS: print("  ✔", m)
if FAIL:
    print("\n  ── 실패 ──")
    for m in FAIL: print("  ✘", m)
else:
    print("\n  ✅ 모든 항목 통과")
