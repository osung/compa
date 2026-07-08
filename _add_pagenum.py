# -*- coding: utf-8 -*-
"""전체 요약부터 새 구역으로 분리 → 페이지번호 1부터, 하단 중앙 표시.
표지/목차 구역은 번호 없음."""
import copy, io, sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = sys.argv[1] if len(sys.argv) > 1 else "COMPA_최종Top5_보고서.docx"
DST = sys.argv[2] if len(sys.argv) > 2 else SRC

d = Document(SRC)

# --- 1) '■ 전체 요약' 바로 앞 문단에 구역 나누기(nextPage) 삽입 ---
paras = d.paragraphs
idx = next(i for i, p in enumerate(paras) if p.text.strip().startswith("■ 전체 요약"))
before = paras[idx - 1]                      # 구역1의 마지막 문단(빈 줄)

body = d.element.body
body_sect = body.find(qn("w:sectPr"))        # 문서 최종 sectPr = 구역2

# 구역1용 sectPr = 최종 sectPr 복제 + type=nextPage (pgNumType 추가 전에 복제)
sect1 = copy.deepcopy(body_sect)
for pn in sect1.findall(qn("w:pgNumType")):  # 혹시 있으면 제거(구역1은 번호 없음)
    sect1.remove(pn)
typ = OxmlElement("w:type"); typ.set(qn("w:val"), "nextPage")
sect1.insert(0, typ)
pPr = before._p.get_or_add_pPr()
pPr.append(sect1)

# --- 2) 구역2(최종 sectPr): 페이지번호 1부터 ---
for pn in body_sect.findall(qn("w:pgNumType")):
    body_sect.remove(pn)
pgnum = OxmlElement("w:pgNumType"); pgnum.set(qn("w:start"), "1")
cols = body_sect.find(qn("w:cols"))          # pgNumType 는 cols 앞에 위치
if cols is not None:
    cols.addprevious(pgnum)
else:
    body_sect.append(pgnum)

# 중간 저장 후 재오픈(구역/푸터 API 안정화)
buf = io.BytesIO(); d.save(buf); buf.seek(0)
d = Document(buf)

# --- 3) 구역2 하단 중앙에 PAGE 필드 ---
secs = d.sections
assert len(secs) == 2, f"구역 수 예상 2, 실제 {len(secs)}"
foot = secs[1].footer
foot.is_linked_to_previous = False
p = foot.paragraphs[0] if foot.paragraphs else foot.add_paragraph()
p.text = ""
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
r = OxmlElement("w:r"); t = OxmlElement("w:t"); t.text = "1"; r.append(t); fld.append(r)
p._p.append(fld)

# 구역1(표지/목차) 푸터는 비움
secs[0].footer.is_linked_to_previous = False
for fp in secs[0].footer.paragraphs:
    fp.text = ""

d.save(DST)
print("페이지번호 추가 완료: 구역", len(secs), "| 전체요약=1p 시작 | saved:", DST)
