# -*- coding: utf-8 -*-
"""각 [수요 N] 아래 기업명·수요기술 내용·수요기술 사양 3문단을 2열 표로 변환.
값 문단은 복제해 줄바꿈/서식 보존. 표 뒤에 매칭 관점 등 기존 내용 유지."""
import copy, sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

SRC = sys.argv[1] if len(sys.argv) > 1 else "COMPA_최종Top5_보고서.docx"
DST = sys.argv[2] if len(sys.argv) > 2 else SRC

LABELS = [("기업명", "기업명:"), ("수요기술 내용", "수요기술 내용:"), ("수요기술 사양", "수요기술 사양:")]
W_LABEL, W_VALUE = 1560, 7080   # dxa (합 8640)

d = Document(SRC)

def set_borders(tbl):
    tblPr = tbl._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        x = OxmlElement(f"w:{e}")
        x.set(qn("w:val"), "single"); x.set(qn("w:sz"), "4")
        x.set(qn("w:space"), "0"); x.set(qn("w:color"), "808080")
        b.append(x)
    tblPr.append(b)
    # 고정 레이아웃 + 전체 너비
    lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"), "fixed"); tblPr.append(lay)
    w = OxmlElement("w:tblW"); w.set(qn("w:w"), str(W_LABEL + W_VALUE)); w.set(qn("w:type"), "dxa")
    tblPr.append(w)

def set_cell_w(cell, w):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
    tcW.set(qn("w:w"), str(w)); tcW.set(qn("w:type"), "dxa")

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def vcenter(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement("w:vAlign"); va.set(qn("w:val"), "center"); tcPr.append(va)

n = 0
for p in list(d.paragraphs):
    if p.style is None or p.style.name != "Heading 2" or not p.text.strip().startswith("[수요"):
        continue
    # 헤딩 다음 형제 문단들에서 3개 라벨 문단 수집
    src = {}
    el = p._p.getnext()
    while el is not None and len(src) < 3:
        if el.tag != qn("w:p"):
            break
        txt = "".join(t.text or "" for t in el.findall(".//" + qn("w:t")))
        matched = False
        for key, pref in LABELS:
            if txt.strip().startswith(pref):
                src[key] = el; matched = True; break
        nxt = el.getnext()
        if not matched and txt.strip() == "":
            el = nxt; continue          # 중간 빈 줄 허용
        if not matched:
            break
        el = nxt
    present = [(label, pref) for (label, pref) in LABELS if label in src]
    if "기업명" not in src or len(present) < 2:
        continue

    # 표 생성(끝에) 후 헤딩 뒤로 이동. 존재하는 필드만 행으로.
    tbl = d.add_table(rows=len(present), cols=2)
    set_borders(tbl)
    for ri, (label, pref) in enumerate(present):
        lc, vc = tbl.rows[ri].cells
        set_cell_w(lc, W_LABEL); set_cell_w(vc, W_VALUE)
        shade(lc, "F2F2F2"); vcenter(lc)
        # 라벨 셀
        lp = lc.paragraphs[0]; lp.text = ""
        run = lp.add_run(label); run.bold = True
        run.font.size = Pt(9)
        # 값 셀: 원본 문단 복제 후 라벨 run 제거
        src_p = copy.deepcopy(src[label])
        # 첫 run(라벨) 제거
        runs = src_p.findall(qn("w:r"))
        if runs:
            # 라벨 접두어를 가진 첫 run 삭제
            first = runs[0]
            first_txt = "".join(t.text or "" for t in first.findall(qn("w:t")))
            if first_txt.strip().startswith(pref.rstrip(":")):
                src_p.remove(first)
        # 값 셀의 기본 빈 문단 제거 후 복제 문단 삽입
        vc._tc.remove(vc._tc.find(qn("w:p")))
        vc._tc.append(src_p)

    p._p.addnext(tbl._tbl)
    # 원본 문단(존재분) 삭제
    for key in list(src):
        src[key].getparent().remove(src[key])
    n += 1

d.save(DST)
print(f"표로 변환한 수요기술: {n}개  (라벨 {W_LABEL} / 값 {W_VALUE} dxa)")
print("saved:", DST)
