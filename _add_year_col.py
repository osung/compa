# -*- coding: utf-8 -*-
"""최종 추천 Top5 표에 '과제수행년도' 열 추가(수행기관 뒤). 값=과제설명문의 시작~종료 연도."""
import copy, re, sys, glob, json, pickle
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = sys.argv[1] if len(sys.argv) > 1 else "COMPA_최종Top5_보고서.docx"
DST = sys.argv[2] if len(sys.argv) > 2 else SRC
SCRATCH = "/private/tmp/claude-501/-Users-osung-work-compa/d6ed121c-12e4-45b4-b2fb-535b7554627c/scratchpad"

def norm(s): return re.sub(r"\s+", "", str(s)).strip()

def period(desc):
    desc = str(desc)
    s = re.search(r"(20\d\d)\s*년\s*\d+\s*월\s*\d+\s*일에\s*시작", desc)
    e = re.search(r"(20\d\d)\s*년\s*\d+\s*월\s*\d+\s*일에\s*종료", desc)
    if s and e:
        sy, ey = s.group(1), e.group(1)
        return sy if sy == ey else f"{sy}~{ey}"
    return None

# --- (과제명,수행기관) -> 수행년도 ---
dmap = {}
for f in glob.glob("COMPA_*_최종추천.pkl"):
    if "전체" in f: continue
    for r in pickle.load(open(f, "rb")).to_dict("records"):
        dmap.setdefault((norm(r["과제명"]), norm(r["과제수행기관"])), str(r.get("과제설명문") or ""))
for k, e in json.load(open("COMPA_통합best.json")).items():
    for t in e.get("top5", []):
        dmap.setdefault((norm(t["과제명"]), norm(t["수행기관"])), str(t.get("과제설명문") or ""))
for t in json.load(open(f"{SCRATCH}/regen_targets.json"))["regen"]:
    dmap.setdefault((norm(t["과제명"]), norm(t["수행기관"])), str(t.get("과제설명문_fallback") or ""))

def year_of(nm, org):
    return period(dmap.get((norm(nm), norm(org)), "")) or ""

WIDTHS = [640, 3360, 1500, 1040, 2100]   # 순위/과제명/수행기관/과제수행년도/매칭근거 (합 8640)
INS = 3   # 수행기관(idx2) 뒤에 삽입

def set_cell_text(cell, value, center=False):
    p = cell.paragraphs[0]
    runs = p.runs
    if not runs:
        p.add_run(value)
    else:
        tgt = max(runs, key=lambda r: len(r.text)) if any(r.text for r in runs) else runs[-1]
        for r in runs:
            r.text = value if r is tgt else ""
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

d = Document(SRC)
n = 0
for tb in d.tables:
    hdr = tuple(c.text.strip() for c in tb.rows[0].cells)
    if hdr[:1] != ("순위",) or "과제수행년도" in hdr:
        continue
    tbl = tb._tbl
    # gridCol 추가
    grid = tbl.find(qn("w:tblGrid"))
    cols = grid.findall(qn("w:gridCol"))
    newg = OxmlElement("w:gridCol")
    cols[INS - 1].addnext(newg)
    for g, w in zip(grid.findall(qn("w:gridCol")), WIDTHS):
        g.set(qn("w:w"), str(w))
    # 각 행에 셀 삽입(같은 행의 수행기관 셀 복제)
    for ri, row in enumerate(tb.rows):
        tcs = row._tr.findall(qn("w:tc"))
        tmpl = copy.deepcopy(tcs[INS - 1])       # 수행기관 셀 복제 → 서식 유지
        tcs[INS - 1].addnext(tmpl)
        # 새 셀 값 설정
        from docx.table import _Cell
        newcell = _Cell(tmpl, tb)
        if ri == 0:
            set_cell_text(newcell, "과제수행년도", center=True)
        else:
            nm = tb.rows[ri].cells[1].text.strip()
            org = tb.rows[ri].cells[2].text.strip()
            set_cell_text(newcell, year_of(nm, org), center=True)
    # 전체 셀 너비 재설정
    for row in tb.rows:
        for cell, w in zip(row.cells, WIDTHS):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
            tcW.set(qn("w:w"), str(w)); tcW.set(qn("w:type"), "dxa")
    n += 1

d.save(DST)
print(f"과제수행년도 열 추가한 Top5 표: {n}개")
print("saved:", DST)
